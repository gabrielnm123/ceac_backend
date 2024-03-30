from django.contrib import admin
from .models import Ficha
from django.http import HttpResponse
from docx import Document

class FichaAdmin(admin.ModelAdmin):
    search_fields = ['nome_completo', 'cpf']
    list_display = ('nome_completo', 'cpf')

    def download_ficha(self, request, queryset):
        document = Document('backend/capacita/doc/FICHA.docx')
        document2 = Document()
        for ficha in queryset:
            for paragraph in document.paragraphs:
                document2.add_paragraph(paragraph.text)
                paragraph.text = paragraph.text.replace('1nome_completo1', ficha.nome_completo)
                paragraph.text = paragraph.text.replace('1cpf1', ficha.cpf)
                if ficha.genero == 'F':
                    paragraph.text = paragraph.text.replace('1f1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1f1', ' ')
                if ficha.genero == 'M':
                    paragraph.text = paragraph.text.replace('1m1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1m1', ' ')
                paragraph.text = paragraph.text.replace('1data_nascimento1', ficha.data_nascimento.strftime('%d/%m/%Y'))
                if ficha.escolaridade == 'fundamental':
                    paragraph.text = paragraph.text.replace('1ensino_fundamental1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1ensino_fundamental1', ' ')
                if ficha.escolaridade == 'medio':
                    paragraph.text = paragraph.text.replace('1ensino_medio1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1ensino_medio1', ' ')
                if ficha.escolaridade == 'graduacao':
                    paragraph.text = paragraph.text.replace('1graduacao1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1graduacao1', ' ')
                if ficha.escolaridade == 'pos_graduacao':
                    paragraph.text = paragraph.text.replace('1pos_graduacao1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1pos_graduacao1', ' ')
                if ficha.area_atuacao == 'artesanato':
                    paragraph.text = paragraph.text.replace('1artesanato1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1artesanato1', ' ')
                if ficha.area_atuacao == 'agricultura':
                    paragraph.text = paragraph.text.replace('1agriculturaU1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1agriculturaU1', ' ')
                if ficha.area_atuacao == 'comercio':
                    paragraph.text = paragraph.text.replace('1comercio1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1comercio1', ' ')
                if ficha.area_atuacao == 'estetica':
                    paragraph.text = paragraph.text.replace('1estet1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1estet1', ' ')
                if ficha.area_atuacao == 'gastronomia':
                    paragraph.text = paragraph.text.replace('1g1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1g1', ' ')
                if ficha.area_atuacao == 'industria':
                    paragraph.text = paragraph.text.replace('1I1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1I1', ' ')
                if ficha.area_atuacao == 'servico':
                    paragraph.text = paragraph.text.replace('1S1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1S1', ' ')
                paragraph.text = paragraph.text.replace('1endereco_residencial1', ficha.endereco)
                if ficha.complemento:
                    paragraph.text = paragraph.text.replace('1complemento1', ficha.complemento)
                else:
                    paragraph.text = paragraph.text.replace('1complemento1', '')
                paragraph.text = paragraph.text.replace('1bairro1', ficha.bairro)
                paragraph.text = paragraph.text.replace('1cep1', ficha.cep)
                paragraph.text = paragraph.text.replace('1uf1', ficha.uf)
                paragraph.text = paragraph.text.replace('1contato1', ficha.contato)
                paragraph.text = paragraph.text.replace('1email1', ficha.email)
                if ficha.interesse_ter_negocio == 's':
                    paragraph.text = paragraph.text.replace('1s_n1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1s_n1', ' ')
                if ficha.interesse_ter_negocio == 'n':
                    paragraph.text = paragraph.text.replace('1n_n1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1n_n1', ' ')
                if ficha.preferencia_aula == 'presencial':
                    paragraph.text = paragraph.text.replace('1presencial1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1presencial1', ' ')
                if ficha.preferencia_aula == 'online':
                    paragraph.text = paragraph.text.replace('1online1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1online1', ' ')
                if ficha.meio_comunicacao_aula == 'whatsapp':
                    paragraph.text = paragraph.text.replace('1zap1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1zap1', ' ')
                if ficha.meio_comunicacao_aula == 'email':
                    paragraph.text = paragraph.text.replace('1email_link1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1email_link1', ' ')
                if ficha.nome_fantasia:
                    paragraph.text = paragraph.text.replace('1nome_fantasia1', ficha.nome_fantasia)
                else:
                    paragraph.text = paragraph.text.replace('1nome_fantasia1', '')
                if ficha.situacao_empresa == 'ativa':
                    paragraph.text = paragraph.text.replace('1ativa1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1ativa1', ' ')
                if ficha.porte_empresa == 'MEI':
                    paragraph.text = paragraph.text.replace('1mei1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1mei1', ' ')
                if ficha.porte_empresa == 'ME':
                    paragraph.text = paragraph.text.replace('1me1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1me1', ' ')
                if ficha.data_abertura:
                    paragraph.text = paragraph.text.replace('1data_abertura1', ficha.data_abertura.strftime('%d/%m/%Y'))
                else:
                    paragraph.text = paragraph.text.replace('1data_abertura1', '')
                if ficha.cnae_principal:
                    paragraph.text = paragraph.text.replace('1cnae1', ficha.cnae_principal)
                else:
                    paragraph.text = paragraph.text.replace('1cnae1', '')
                if ficha.setor == 'comercio':
                    paragraph.text = paragraph.text.replace('1comercio1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1comercio1', ' ')
                if ficha.setor == 'servico':
                    paragraph.text = paragraph.text.replace('1servico1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1servico1', ' ')
                if ficha.setor == 'agronegocios':
                    paragraph.text = paragraph.text.replace('1agro1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1agro1', ' ')
                if ficha.setor == 'industria':
                    paragraph.text = paragraph.text.replace('1indus1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1indus1', ' ')
                if ficha.tipo_vinculo == 'representante':
                    paragraph.text = paragraph.text.replace('1repre1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1repre1', ' ')
                if ficha.tipo_vinculo == 'responsavel':
                    paragraph.text = paragraph.text.replace('1resp1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1resp1', ' ')
                if ficha.assistir_online == 's':
                    paragraph.text = paragraph.text.replace('1s_o1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1s_o1', ' ')
                if ficha.assistir_online == 'n':
                    paragraph.text = paragraph.text.replace('1n_o1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1n_o1', ' ')
                if ficha.if_true_assistir_casa == 'computador':
                    paragraph.text = paragraph.text.replace('1pc1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1pc1', ' ')
                if ficha.if_true_assistir_casa == 'celular':
                    paragraph.text = paragraph.text.replace('1cel1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1cel1', ' ')
                if ficha.if_true_assistir_casa == 'tablet':
                    paragraph.text = paragraph.text.replace('1tablet1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1tablet1', ' ')
                if ficha.if_true_assistir_casa == 'outro':
                    paragraph.text = paragraph.text.replace('1outro1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1outro1', ' ')
                if ficha.modulo_marketing:
                    paragraph.text = paragraph.text.replace('1marketing1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1marketing1', ' ')
                if ficha.modulo_financeiro:
                    paragraph.text = paragraph.text.replace('1financeiro1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1financeiro1', ' ')
                if ficha.modulo_planejamento:
                    paragraph.text = paragraph.text.replace('1planejamento1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1planejamento1', ' ')
                if ficha.modulo_outros:
                    paragraph.text = paragraph.text.replace('1outros1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1outros1', ' ')
                if ficha.responsabilizacao:
                    paragraph.text = paragraph.text.replace('1responsabilizacao1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1responsabilizacao1', ' ')
                if ficha.manejo_dados:
                    paragraph.text = paragraph.text.replace('1manejo_dados1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1manejo_dados1', ' ')
                if ficha.armazenamento_dados:
                    paragraph.text = paragraph.text.replace('1armazenamento_dados1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1armazenamento_dados1', ' ')
                if ficha.autorizacao:
                    paragraph.text = paragraph.text.replace('1autorizacao1', 'x')
                else:
                    paragraph.text = paragraph.text.replace('1autorizacao1', ' ')

        # response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        # response['Content-Disposition'] = 'attachment; filename="ficha.docx"'
        # document.save(response)
        # return response
        document.save('backend/capacita/doc/exemplo.docx')
        document2.save('backend/capacita/doc/exemplo2.docx')


admin.site.register(Ficha, FichaAdmin)
